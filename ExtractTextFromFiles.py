import PyPDF2
import docx
import re

def getPDFContentAsString(filename):
    file = open(filename, 'rb')
    pdfFile = PyPDF2.PdfFileReader(file)
    pageCnt = pdfFile.getNumPages()
    pdfText = ""
    for index in range(pageCnt):
        page = pdfFile.getPage(index)
        pdfText += page.extractText()
    file.close()
    return pdfText.lower()

def getTxtContentAsString(filename):
    jdFile = open(filename, 'r')
    jobDescriptionText = jdFile.read()
    jobDescriptionText = " ".join(jobDescriptionText.replace("•", "").replace("\n", "\t").replace(")","").split())
    jdFile.close()
    return jobDescriptionText

def getDocXContentAsString(filename):
    jDDoc = docx.Document(filename)
    fullText = ''

    for para in jDDoc.paragraphs:
        fullText = fullText + ' ' + para.text

    tbl = jDDoc.tables
    for table in jDDoc.tables:
        for row in table.rows:
            rowText = ''
            for cell in row.cells:
                rowText = rowText + ' ' + cell.text
            fullText = fullText + ' '+ rowText
    fullText = re.sub('\W+', ' ', fullText) 
    fullText = re.sub('[^A-Za-z]+', ' ', fullText)
    fullText = fullText.lower()

    return fullText